import os
import xarray as xr
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from scipy import stats
import matplotlib.ticker as mticker
from matplotlib.lines import Line2D
import regionmask

# Word 三线表导出依赖：pip install python-docx
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================
# 1. 基本参数
# ============================================================
CROPS = ["maize", "rice", "soybeans", "wheat"]

CROP_TITLES = {
    "maize": "Maize",
    "rice": "Rice",
    "soybeans": "Soybean",
    "wheat": "Wheat"
}

# 同时处理 DFAI = 1 和 DFAI = -1
# 绘图时 1 在上面，-1 在下面
DFAI_LIST = [1, -1]

# Word 表格和图中显示名称：DFAI = 1 对应 DTF，DFAI = -1 对应 FTD。
DFAI_DISPLAY_NAMES = {
    1: "DTF",
    -1: "FTD"
}

START_YEAR = 1982
END_YEAR = 2016
Data_from = 'CSIC'
DATA_ROOT = r"C:\Users\lxh\Desktop\NongZuoWu_ChanLiang\New_result\Data\Years_figure17_\CSIC"
OUT_DIR = r"D:\Python\DssAat\NO_2\Figure3\combined_DFAI_1_minus1_4row_4col_CSIC"
os.makedirs(OUT_DIR, exist_ok=True)

MIN_VALID_YEARS = 10

# 保存国家尺度 CSV / 绘图数据时，只保留有效栅格点数 > 10 的记录。
# 这里的“有效栅格点”定义为：该栅格在 1982–2016 年期间至少有 MIN_VALID_YEARS 个有效值。
MIN_VALID_GRID_CELLS = 10

# 生成 Word 三线表时，国家尺度记录使用更严格筛选：只保留有效栅格点数 > 50 的国家。
# 注意：该阈值只用于三线表，不影响前面国家趋势 CSV 和图。
THREE_LINE_TABLE_MIN_VALID_GRID_CELLS = 50

# 第二行/第四行国家柱状图绘制哪个指标
# slope_per_year   ：每年变化率
# slope_per_decade ：每 10 年变化量
PLOT_VALUE = "slope_per_decade"

# 全球折线图阶段性趋势
PERIODS = [
    (1982, 1995),
    (1996, 2005),
    (2006, 2016)
]

VARIABLES = [
    "plant_residual",
    "growth_residual",
    "harvest_residual"
]

LABELS = {
    "plant_residual": "Planting",
    "growth_residual": "Growth",
    "harvest_residual": "Harvest"
}

COLORS = {
    "plant_residual": "#0072B2",
    "growth_residual": "#D55E00",
    "harvest_residual": "#009E73",

    "Planting": "#0072B2",
    "Growth": "#D55E00",
    "Harvest": "#009E73"
}

DESIRED_COLS = ["Planting", "Growth", "Harvest"]

# Word 三线表最终只保留这些列
THREE_LINE_TABLE_COLS = [
    "table_group",
    "region",
    "abbrev",
    "label",
    "slope_per_year",
    "slope_per_decade",
    "r_value",
    "r_squared",
    "p_value",
    "std_err",
    "mean_residual_1982_2016",
    "std_residual_1982_2016"
]

LAT_NAME = "lat"
LON_NAME = "lon"

# ============================================================
# 2. Nature-like 绘图风格
# ============================================================
plt.rcParams.update({
    "font.family": "Arial",
    "font.sans-serif": ["Arial", "SimHei"],
    "font.size": 7,
    "axes.linewidth": 0.6,
    "axes.labelsize": 8,
    "axes.titlesize": 9,
    "xtick.labelsize": 7,
    "ytick.labelsize": 7,
    "legend.fontsize": 7,

    "xtick.direction": "out",
    "ytick.direction": "out",
    "xtick.major.size": 2.5,
    "ytick.major.size": 2.5,
    "xtick.major.width": 0.6,
    "ytick.major.width": 0.6,

    "axes.unicode_minus": False,
    "figure.dpi": 300,
    "savefig.dpi": 600,
    "pdf.fonttype": 42,
    "ps.fonttype": 42
})

# ============================================================
# 3. 国家列表：国家趋势图使用
# CSV 仍保存所有国家，图上只显示这里指定国家
# ============================================================
crop_country_order = {
    "maize": [
        "United States of America",
        "China",
        "Brazil",
        "Argentina",
        "Mexico"
    ],

    "rice": [
        "China",
        "India",
        "Bangladesh",
        "Indonesia",
        "Vietnam"
    ],

    "soybeans": [
        "Brazil",
        "United States of America",
        "Argentina",
        "China",
        "India"
    ],

    "wheat": [
        "China",
        "United States of America",
        "Ukraine",
        "Australia",
        "Canada",
    ]
}

country_aliases = {
    "United States of America": [
        "United States of America",
        "United States",
        "USA"
    ],
    "Russia": [
        "Russia",
        "Russian Federation"
    ],
    "United Kingdom": [
        "United Kingdom",
        "UK",
        "Great Britain"
    ],
    "Czechia": [
        "Czechia",
        "Czech Republic"
    ],
    "South Korea": [
        "South Korea",
        "Republic of Korea",
        "Korea, Republic of"
    ],
    "Iran": [
        "Iran",
        "Iran, Islamic Republic of"
    ],
    "Vietnam": [
        "Vietnam",
        "Viet Nam"
    ],
    "Laos": [
        "Laos",
        "Lao PDR",
        "Lao People's Democratic Republic"
    ],
    "Tanzania": [
        "Tanzania",
        "United Republic of Tanzania"
    ],
    "Bolivia": [
        "Bolivia",
        "Bolivia, Plurinational State of"
    ],
    "Myanmar": [
        "Myanmar",
        "Burma"
    ],
    "Guinea": [
        "Guinea"
    ],
    "Serbia": [
        "Serbia"
    ]
}

country_display_names = {
    "United States of America": "USA"
}

# ============================================================
# 4. 工具函数
# ============================================================
def normalize_country_name(name):
    return (
        str(name)
        .lower()
        .replace(".", "")
        .replace(",", "")
        .replace("'", "")
        .replace("’", "")
        .strip()
    )


def get_selected_country_info(country_name, target_countries):
    """
    判断 Natural Earth 国家名是否属于目标国家列表。
    """
    country_order_dict = {
        country: i for i, country in enumerate(target_countries)
    }

    country_name_norm = normalize_country_name(country_name)

    for target_country in target_countries:
        candidate_names = [target_country]

        if target_country in country_aliases:
            candidate_names += country_aliases[target_country]

        candidate_norms = [
            normalize_country_name(x) for x in candidate_names
        ]

        if country_name_norm in candidate_norms:
            return True, target_country, country_order_dict.get(target_country, 9999)

    return False, country_name, 9999


def get_country_regions():
    """
    兼容不同 regionmask 版本。
    """
    try:
        return regionmask.defined_regions.natural_earth_v5_0_0.countries_110
    except AttributeError:
        try:
            return regionmask.defined_regions.natural_earth_v5_1_2.countries_110
        except AttributeError:
            return regionmask.defined_regions.natural_earth_v4_1_0.countries_110


def make_country_mask(countries, ds_use):
    """
    构建国家掩膜。
    """
    try:
        country_mask = countries.mask(
            ds_use,
            lon_name=LON_NAME,
            lat_name=LAT_NAME
        )
    except TypeError:
        country_mask = countries.mask(
            ds_use[LON_NAME],
            ds_use[LAT_NAME]
        )

    return country_mask


def linear_fit_ci(x, y, confidence=0.95):
    """
    计算线性拟合线及其 95% 置信区间。
    返回：
    x_fit, y_fit, ci_lower, ci_upper
    """
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)

    mask = np.isfinite(y)
    x = x[mask]
    y = y[mask]

    n = len(x)

    if n <= 2:
        return None, None, None, None

    slope, intercept = np.polyfit(x, y, 1)

    x_fit = np.linspace(x.min(), x.max(), 200)
    y_fit = slope * x_fit + intercept

    y_pred = slope * x + intercept
    residuals = y - y_pred

    dof = n - 2
    residual_std_error = np.sqrt(np.sum(residuals ** 2) / dof)

    t_value = stats.t.ppf((1 + confidence) / 2, dof)

    x_mean = np.mean(x)
    sxx = np.sum((x - x_mean) ** 2)

    fit_se = residual_std_error * np.sqrt(
        1 / n + (x_fit - x_mean) ** 2 / sxx
    )

    ci = t_value * fit_se

    ci_lower = y_fit - ci
    ci_upper = y_fit + ci

    return x_fit, y_fit, ci_lower, ci_upper


def set_symmetric_ylim(ax, values, nbins=4):
    """
    全球折线图每个子图单独根据自身数据设置 y 轴范围。
    """
    values = np.asarray(values, dtype=float)
    values = values[np.isfinite(values)]

    if values.size == 0:
        return

    max_abs = np.nanmax(np.abs(values))

    if max_abs == 0:
        max_abs = 1

    lim = max_abs * 1.18
    ax.set_ylim(-lim, lim)
    ax.yaxis.set_major_locator(mticker.MaxNLocator(nbins=nbins))


def set_symmetric_xlim(ax, values, nbins=4):
    """
    国家/全球趋势图每个子图单独根据趋势值设置 x 轴范围。
    """
    values = np.asarray(values, dtype=float)
    values = values[np.isfinite(values)]

    if values.size == 0:
        return

    max_abs = np.nanmax(np.abs(values))

    if max_abs == 0:
        max_abs = 1

    lim = max_abs * 1.18
    ax.set_xlim(-lim, lim)
    ax.xaxis.set_major_locator(mticker.MaxNLocator(nbins=nbins))


def style_axis(ax):
    """
    Nature-like 坐标轴样式。
    """
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    ax.spines["left"].set_linewidth(0.6)
    ax.spines["bottom"].set_linewidth(0.6)

    ax.tick_params(
        axis="both",
        which="major",
        direction="out",
        length=2.5,
        width=0.6,
        pad=2
    )

    ax.grid(False)



def count_valid_grid_cells(da, min_valid_years=MIN_VALID_YEARS):
    """
    统计有效栅格点数。

    有效栅格点定义：在 START_YEAR–END_YEAR 时段内，
    该栅格至少有 min_valid_years 个非 NaN 年份。

    返回：
    valid_grid_mask, n_valid_grid_cells
    """
    da_period = da.where(
        (da["year"] >= START_YEAR) & (da["year"] <= END_YEAR),
        drop=True
    )

    valid_grid_mask = da_period.notnull().sum(dim="year") >= min_valid_years
    n_valid_grid_cells = int(valid_grid_mask.sum().values)

    return valid_grid_mask, n_valid_grid_cells


# ============================================================
# 4b. Word 三线表工具函数：Nature-like table style
# ============================================================
def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """
    设置单元格边框。val='nil' 表示无边框；val='single' 表示实线。
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")

    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    border_map = {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right
    }

    for edge, border in border_map.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)

        if border is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), border.get("val", "single"))
            element.set(qn("w:sz"), str(border.get("sz", 6)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), border.get("color", "000000"))


def set_cell_text(cell, text, bold=False, font_size=7):
    """
    设置单元格文字，统一 Times New Roman，小字号，适合期刊三线表。
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table_value(value):
    """
    Word 表格数值格式：过长小数保留 3 位，有效位更清晰。
    """
    if pd.isna(value):
        return ""

    if isinstance(value, (float, np.floating)):
        return f"{value:.3g}"

    return value


def make_nature_three_line_table(document, df, caption, font_size=7):
    """
    在 Word 中写入 Nature-like 三线表：
    - 无竖线；
    - 表头上方一条粗线；
    - 表头下方一条细线；
    - 表格底部一条粗线。
    """
    if df is None or df.empty:
        document.add_paragraph(f"{caption}: No data.")
        return

    caption_paragraph = document.add_paragraph()
    caption_run = caption_paragraph.add_run(caption)
    caption_run.bold = True
    caption_run.font.size = Pt(8)
    caption_run.font.name = "Times New Roman"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption_paragraph.paragraph_format.space_after = Pt(3)

    df_write = df.copy()

    # 保存 Word 三线表时只保留指定列。
    keep_cols = [c for c in THREE_LINE_TABLE_COLS if c in df_write.columns]

    if len(keep_cols) > 0:
        df_write = df_write[keep_cols]

    table = document.add_table(rows=1, cols=len(df_write.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header_cells = table.rows[0].cells
    for j, col in enumerate(df_write.columns):
        set_cell_text(header_cells[j], col, bold=True, font_size=font_size)

    for _, row in df_write.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df_write.columns):
            set_cell_text(cells[j], format_table_value(row[col]), font_size=font_size)

    # 先清除所有边框，避免出现竖线和网格线。
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=None, bottom=None, left=None, right=None)

    top_border = {"val": "single", "sz": 12, "color": "000000"}
    mid_border = {"val": "single", "sz": 6, "color": "000000"}
    bottom_border = {"val": "single", "sz": 12, "color": "000000"}

    # 第一线：表头上边框；第二线：表头下边框。
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=top_border, bottom=mid_border, left=None, right=None)

    # 第三线：最后一行下边框。
    for cell in table.rows[-1].cells:
        set_cell_border(cell, top=None, bottom=bottom_border, left=None, right=None)

    document.add_paragraph()


def save_tables_to_nature_word(tables, out_docx, title=None):
    """
    将多个 DataFrame 写入同一个 Word 文件，格式为 Nature-like 三线表。

    Parameters
    ----------
    tables : list of tuple
        [(caption, dataframe), ...]
    out_docx : str
        Word 输出路径。
    title : str, optional
        文档标题。
    """
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(7)

    if title is not None:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(10)
        title_run.font.name = "Times New Roman"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        document.add_paragraph()

    for caption, df in tables:
        make_nature_three_line_table(document, df, caption)

    document.save(out_docx)
    print("Nature-like 三线表 Word 已保存：", out_docx)


def build_crop_dfai_trend_table(crop, dfai, global_trend_df, country_trend_df):
    """
    为 Word 输出构建单个作物 × 单个 DTF/FTD 的三线表数据。

    每张表对应一个作物和一个 DFAI 类型：
    4 个作物 × 2 个类型 = 8 张表。

    表内合并 Global 整体趋势和国家尺度趋势；生成 Word 三线表时，
    国家尺度记录进一步按 n_valid_grid_cells > THREE_LINE_TABLE_MIN_VALID_GRID_CELLS 过滤。
    """
    dfs = []

    if global_trend_df is not None and not global_trend_df.empty:
        global_part = global_trend_df[
            (global_trend_df["crop"] == crop) &
            (global_trend_df["DFAI"] == dfai)
        ].copy()
        if not global_part.empty:
            dfs.append(global_part)

    if country_trend_df is not None and not country_trend_df.empty:
        country_part = country_trend_df[
            (country_trend_df["crop"] == crop) &
            (country_trend_df["DFAI"] == dfai)
        ].copy()

        # 只在生成 Word 三线表时使用更严格阈值：
        # 国家尺度记录仅保留有效数据栅格数量 > 50 的国家；Global 记录不受此处影响。
        if not country_part.empty and "n_valid_grid_cells" in country_part.columns:
            country_part = country_part[
                country_part["n_valid_grid_cells"] > THREE_LINE_TABLE_MIN_VALID_GRID_CELLS
            ].copy()

        if not country_part.empty:
            dfs.append(country_part)

    if len(dfs) == 0:
        return pd.DataFrame()

    table_df = pd.concat(dfs, ignore_index=True, sort=False)

    dfai_label = DFAI_DISPLAY_NAMES.get(dfai, str(dfai))
    crop_name = CROP_TITLES.get(crop, crop)

    table_df["DFAI_label"] = dfai_label
    table_df["crop_name"] = crop_name
    table_df["table_group"] = f"{crop_name}-{dfai_label}"

    # Word 表格中显示的地区名：Global 保留 Global；目标国家用 selected_country_name；其他国家用 Natural Earth 名称。
    table_df["region"] = np.where(
        table_df.get("is_global", False).astype(bool),
        "Global",
        table_df.get("selected_country_name", "").replace("", np.nan)
    )
    table_df["region"] = table_df["region"].fillna(table_df["country"])

    # Global 排在第一位；被选中的重点国家按 crop_country_order 排序；其他国家按名称排序。
    table_df["_is_global_sort"] = table_df.get("is_global", False).astype(bool).astype(int)
    if "country_order" not in table_df.columns:
        table_df["country_order"] = 9999
    table_df["country_order"] = table_df["country_order"].fillna(9999)

    label_order = {label: i for i, label in enumerate(DESIRED_COLS)}
    table_df["_label_order"] = table_df["label"].map(label_order).fillna(9999)

    table_df = table_df.sort_values(
        by=["_is_global_sort", "country_order", "region", "_label_order"],
        ascending=[False, True, True, True]
    )

    output_cols = [col for col in THREE_LINE_TABLE_COLS if col in table_df.columns]
    return table_df[output_cols].reset_index(drop=True)


def build_eight_crop_dfai_tables(all_global_trend_df, all_country_df):
    """
    生成 8 张 Word 三线表：
    DTF 下 4 个作物 + FTD 下 4 个作物。
    """
    tables = []
    table_id = 1

    # 按 DFAI 分组输出：先 DTF 的四个作物，再 FTD 的四个作物。
    for dfai in DFAI_LIST:
        dfai_label = DFAI_DISPLAY_NAMES.get(dfai, str(dfai))

        for crop in CROPS:
            crop_name = CROP_TITLES.get(crop, crop)
            table_df = build_crop_dfai_trend_table(
                crop=crop,
                dfai=dfai,
                global_trend_df=all_global_trend_df,
                country_trend_df=all_country_df
            )

            caption = (
                f"Table {table_id} | {crop_name} {dfai_label} trends, "
                f"{START_YEAR}–{END_YEAR}. Country-level records in this Word table retain only "
                f"regions with more than {THREE_LINE_TABLE_MIN_VALID_GRID_CELLS} valid grid cells."
            )

            tables.append((caption, table_df))
            table_id += 1

    return tables


# ============================================================
# 5. 计算单个作物、单个 DFAI 的全球均值、阶段趋势、全球整体趋势、国家趋势
# ============================================================
def process_crop(crop, dfai, country_regions):
    print(f"\n正在处理：crop = {crop}, DFAI = {dfai}")

    file_path = fr"{DATA_ROOT}\{crop}\ls_{dfai}_multimodel_aic_yearly_residuals_-60_60_-180_180.nc"

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在：{file_path}")

    ds = xr.open_dataset(file_path)

    ds_use = ds.where(
        (ds["year"] >= START_YEAR) & (ds["year"] <= END_YEAR),
        drop=True
    )

    years = ds_use["year"].values.astype(float)

    weights = np.cos(np.deg2rad(ds_use[LAT_NAME]))
    weights.name = "weights"

    # ------------------------------------------------------------
    # 5.1 全球面积加权平均
    # ------------------------------------------------------------
    global_dict = {
        "year": years.astype(int)
    }

    for var in VARIABLES:
        da = ds_use[var]

        global_mean = da.weighted(weights).mean(
            dim=[LAT_NAME, LON_NAME],
            skipna=True
        )

        global_dict[var] = global_mean.values.astype(float)

    global_df = pd.DataFrame(global_dict)
    global_df["crop"] = crop
    global_df["DFAI"] = dfai

    out_global_csv = fr"{OUT_DIR}\{crop}_global_mean_residual_{START_YEAR}_{END_YEAR}_DFAI_{dfai}_{Data_from}.csv"
    global_df.to_csv(out_global_csv, index=False, encoding="utf-8-sig")
    print("全球均值 CSV 已保存：", out_global_csv)

    # ------------------------------------------------------------
    # 5.2 全球阶段趋势统计
    # ------------------------------------------------------------
    stage_records = []

    for var in VARIABLES:
        y = global_df[var].values.astype(float)

        for start_year, end_year in PERIODS:
            mask = (
                (years >= start_year) &
                (years <= end_year) &
                np.isfinite(y)
            )

            if mask.sum() <= 2:
                continue

            x_stage = years[mask]
            y_stage = y[mask]

            slope, intercept, r_value, p_value, std_err = stats.linregress(
                x_stage,
                y_stage
            )

            stage_records.append({
                "DFAI": dfai,
                "crop": crop,
                "variable": var,
                "label": LABELS[var],
                "period": f"{start_year}-{end_year}",
                "slope_per_year": slope,
                "slope_per_decade": slope * 10,
                "intercept": intercept,
                "r_value": r_value,
                "r_squared": r_value ** 2,
                "p_value": p_value,
                "std_err": std_err,
                "n": int(mask.sum())
            })

    stage_df = pd.DataFrame(stage_records)

    out_stage_csv = fr"{OUT_DIR}\{crop}_global_stage_trend_{START_YEAR}_{END_YEAR}_DFAI_{dfai}_{Data_from}.csv"
    stage_df.to_csv(out_stage_csv, index=False, encoding="utf-8-sig")
    print("全球阶段趋势 CSV 已保存：", out_stage_csv)

    # ------------------------------------------------------------
    # 5.2b 全球整体趋势：用于加入国家柱状图
    #      计算方法与国家趋势完全一致：1982–2016 整段线性趋势
    # ------------------------------------------------------------
    global_trend_records = []

    for var in VARIABLES:
        # 全局有效栅格点数：用于记录到 CSV / Word 表中。
        valid_grid_mask, n_valid_grid_cells = count_valid_grid_cells(
            ds_use[var],
            min_valid_years=MIN_VALID_YEARS
        )

        if n_valid_grid_cells <= MIN_VALID_GRID_CELLS:
            continue

        y = global_df[var].values.astype(float)

        valid_mask = (
            (years >= START_YEAR) &
            (years <= END_YEAR) &
            np.isfinite(y)
        )

        if valid_mask.sum() < MIN_VALID_YEARS:
            continue

        x_valid = years[valid_mask]
        y_valid = y[valid_mask]

        slope, intercept, r_value, p_value, std_err = stats.linregress(
            x_valid,
            y_valid
        )

        global_trend_records.append({
            "DFAI": dfai,
            "crop": crop,
            "country": "Global",
            "selected_country_name": "Global",
            "abbrev": "GLB",
            "is_selected_country": True,
            "country_order": -1,       # 让 Global 排在最上面
            "is_global": True,
            "variable": var,
            "label": LABELS[var],

            "slope_per_year": slope,
            "slope_per_decade": slope * 10,

            "intercept": intercept,
            "r_value": r_value,
            "r_squared": r_value ** 2,
            "p_value": p_value,
            "std_err": std_err,

            "n_valid_years": int(valid_mask.sum()),
            "n_valid_grid_cells": n_valid_grid_cells,
            "first_year": int(x_valid.min()),
            "last_year": int(x_valid.max()),
            "mean_residual_1982_2016": np.nanmean(y_valid),
            "std_residual_1982_2016": np.nanstd(y_valid, ddof=1)
        })

    global_trend_df = pd.DataFrame(global_trend_records)

    out_global_trend_csv = fr"{OUT_DIR}\{crop}_global_overall_trend_{START_YEAR}_{END_YEAR}_DFAI_{dfai}_{Data_from}.csv"
    global_trend_df.to_csv(out_global_trend_csv, index=False, encoding="utf-8-sig")
    print("全球整体趋势 CSV 已保存：", out_global_trend_csv)

    # ------------------------------------------------------------
    # 5.3 国家尺度趋势
    # ------------------------------------------------------------
    country_mask = make_country_mask(country_regions, ds_use)

    if crop not in crop_country_order:
        raise ValueError(f"{crop} 没有设置目标国家列表。")

    target_countries = crop_country_order[crop]

    all_country_records = []

    for country_number, country_name, country_abbrev in zip(
        country_regions.numbers,
        country_regions.names,
        country_regions.abbrevs
    ):
        country_bool = country_mask == country_number

        if int(country_bool.sum().values) == 0:
            continue

        is_selected, matched_country, country_order = get_selected_country_info(
            country_name,
            target_countries
        )

        for var in VARIABLES:
            da = ds_use[var]

            da_country = da.where(country_bool)

            # 仅保留有效栅格点数 > MIN_VALID_GRID_CELLS 的国家-变量记录。
            # 有效栅格点：1982–2016 年期间至少有 MIN_VALID_YEARS 个有效年份。
            valid_grid_mask, n_valid_grid_cells = count_valid_grid_cells(
                da_country,
                min_valid_years=MIN_VALID_YEARS
            )

            if n_valid_grid_cells <= MIN_VALID_GRID_CELLS:
                continue

            # 国家年均值也只基于这些有效栅格点计算，保证 CSV / Word / 绘图一致。
            da_country = da_country.where(valid_grid_mask)

            country_mean_yearly = da_country.weighted(weights).mean(
                dim=[LAT_NAME, LON_NAME],
                skipna=True
            )

            y = country_mean_yearly.values.astype(float)

            valid_mask = (
                (years >= START_YEAR) &
                (years <= END_YEAR) &
                np.isfinite(y)
            )

            if valid_mask.sum() < MIN_VALID_YEARS:
                continue

            x_valid = years[valid_mask]
            y_valid = y[valid_mask]

            slope, intercept, r_value, p_value, std_err = stats.linregress(
                x_valid,
                y_valid
            )

            all_country_records.append({
                "DFAI": dfai,
                "crop": crop,
                "country": country_name,
                "selected_country_name": matched_country if is_selected else "",
                "abbrev": country_abbrev,
                "is_selected_country": is_selected,
                "country_order": country_order,
                "is_global": False,
                "variable": var,
                "label": LABELS[var],

                "slope_per_year": slope,
                "slope_per_decade": slope * 10,

                "intercept": intercept,
                "r_value": r_value,
                "r_squared": r_value ** 2,
                "p_value": p_value,
                "std_err": std_err,

                "n_valid_years": int(valid_mask.sum()),
                "n_valid_grid_cells": n_valid_grid_cells,
                "first_year": int(x_valid.min()),
                "last_year": int(x_valid.max()),
                "mean_residual_1982_2016": np.nanmean(y_valid),
                "std_residual_1982_2016": np.nanstd(y_valid, ddof=1)
            })

    all_country_df = pd.DataFrame(all_country_records)

    # 双重保险：保存 CSV 前再次确保只保留有效栅格点数 > 10 的记录。
    if not all_country_df.empty and "n_valid_grid_cells" in all_country_df.columns:
        all_country_df = all_country_df[
            all_country_df["n_valid_grid_cells"] > MIN_VALID_GRID_CELLS
        ].copy()

    out_country_csv = fr"{OUT_DIR}\{crop}_all_country_trend_{START_YEAR}_{END_YEAR}_DFAI_{dfai}_{Data_from}.csv"
    all_country_df.to_csv(out_country_csv, index=False, encoding="utf-8-sig")
    print("国家趋势 CSV 已保存：", out_country_csv)

    # ------------------------------------------------------------
    # 5.4 组织柱状图数据：Global + 目标国家
    # ------------------------------------------------------------
    if all_country_df.empty:
        country_plot_source_df = pd.DataFrame()
    else:
        country_plot_source_df = all_country_df[
            all_country_df["is_selected_country"] == True
        ].copy()

        country_plot_source_df["plot_country"] = country_plot_source_df["selected_country_name"]
        country_plot_source_df["is_global"] = False

        matched_countries = set(country_plot_source_df["plot_country"].unique())

        missing_countries = [
            c for c in target_countries if c not in matched_countries
        ]

        if len(missing_countries) > 0:
            print(f"{crop}, DFAI = {dfai} 以下目标国家没有匹配到或没有有效数据：")
            for c in missing_countries:
                print("  ", c)

    # Global 作为柱状图第一行
    if not global_trend_df.empty:
        global_plot_source_df = global_trend_df.copy()
        global_plot_source_df["plot_country"] = "Global"
        global_plot_source_df["is_global"] = True
    else:
        global_plot_source_df = pd.DataFrame()

    plot_source_df = pd.concat(
        [global_plot_source_df, country_plot_source_df],
        ignore_index=True,
        sort=False
    )

    if plot_source_df.empty:
        plot_wide = pd.DataFrame(
            columns=["plot_country", "country_order", "is_global"] + DESIRED_COLS
        )
    else:
        plot_wide = plot_source_df.pivot_table(
            index=["plot_country", "country_order", "is_global"],
            columns="label",
            values=PLOT_VALUE,
            aggfunc="mean"
        ).reset_index()

        plot_wide = plot_wide.reindex(
            columns=["plot_country", "country_order", "is_global"] + DESIRED_COLS
        )

        plot_wide["is_global"] = plot_wide["is_global"].fillna(False)

        plot_wide = plot_wide.sort_values(
            "country_order",
            ascending=True
        )

    ds.close()

    return global_df, stage_df, all_country_df, global_trend_df, plot_wide


# ============================================================
# 6. 绘制全球折线图
# ============================================================
def plot_global_panel(ax, crop, global_df, panel_label, show_ylabel=False):
    years = global_df["year"].values.astype(float)

    all_y_for_ylim = []

    for var in VARIABLES:
        y = global_df[var].values.astype(float)

        all_y_for_ylim.append(y)

        # 原始逐年折线
        ax.plot(
            years,
            y,
            color=COLORS[var],
            linewidth=0.50,
            alpha=0.88
        )

        # 阶段性拟合线 + 95% CI
        for start_year, end_year in PERIODS:
            mask = (
                (years >= start_year) &
                (years <= end_year) &
                np.isfinite(y)
            )

            x_stage = years[mask]
            y_stage = y[mask]

            if len(x_stage) <= 2:
                continue

            x_fit, y_fit, ci_lower, ci_upper = linear_fit_ci(
                x_stage,
                y_stage,
                confidence=0.95
            )

            if x_fit is None:
                continue

            ax.fill_between(
                x_fit,
                ci_lower,
                ci_upper,
                color=COLORS[var],
                alpha=0.07,
                linewidth=0
            )

            ax.plot(
                x_fit,
                y_fit,
                color=COLORS[var],
                linestyle="--",
                linewidth=0.70,
                alpha=0.95
            )

    # 阶段分界线
    for boundary in [1995.5, 2005.5]:
        ax.axvline(
            boundary,
            color="black",
            linestyle=":",
            linewidth=0.45,
            alpha=0.35
        )

    # 零线
    ax.axhline(
        0,
        color="black",
        linestyle="--",
        linewidth=0.45,
        alpha=0.6
    )

    ax.set_xlim(START_YEAR, END_YEAR)
    ax.set_xticks([1985, 1995, 2005, 2015])

    all_y_for_ylim = np.concatenate(all_y_for_ylim)
    set_symmetric_ylim(ax, all_y_for_ylim, nbins=4)

    if show_ylabel:
        ax.set_ylabel("Global mean\nyield residual")
    else:
        ax.set_ylabel("")

    style_axis(ax)

    ax.text(
        -0.10,
        1.03,
        panel_label,
        transform=ax.transAxes,
        fontsize=9,
        fontweight="bold",
        va="bottom",
        ha="left"
    )


# ============================================================
# 7. 绘制 Global + 国家趋势条形图
# ============================================================
def plot_country_panel(ax, crop, plot_wide, panel_label, show_ylabel=False):
    if plot_wide.empty:
        ax.text(
            0.5,
            0.5,
            "No country data",
            ha="center",
            va="center",
            transform=ax.transAxes,
            fontsize=7
        )
        ax.set_xticks([])
        ax.set_yticks([])
        style_axis(ax)
        return

    countries_plot = plot_wide["plot_country"].tolist()
    countries_plot_display = [
        country_display_names.get(c, c) for c in countries_plot
    ]

    y_pos = np.arange(len(countries_plot))

    if "is_global" in plot_wide.columns:
        is_global_row = plot_wide["is_global"].astype(bool).values
    else:
        is_global_row = np.zeros(len(plot_wide), dtype=bool)

    bar_height = 0.19

    offsets = {
        "Planting": -bar_height,
        "Growth": 0,
        "Harvest": bar_height
    }

    # ------------------------------------------------------------
    # 突出 Global 行：浅灰背景 + 分隔线
    # ------------------------------------------------------------
    if np.any(is_global_row):
        global_indices = np.where(is_global_row)[0]

        for gi in global_indices:
            ax.axhspan(
                gi - 0.48,
                gi + 0.48,
                facecolor="0.92",
                edgecolor="0.35",
                linewidth=0.55,
                zorder=0
            )

            ax.axhline(
                gi + 0.55,
                color="0.35",
                linestyle=":",
                linewidth=0.55,
                zorder=1
            )

    # ------------------------------------------------------------
    # 普通国家：无边框
    # Global：黑色边框 + hatch，和国家区分
    # ------------------------------------------------------------
    normal_mask = ~is_global_row
    global_mask = is_global_row

    for label in DESIRED_COLS:
        values = plot_wide[label].values.astype(float)

        # 国家柱
        if np.any(normal_mask):
            ax.barh(
                y_pos[normal_mask] + offsets[label],
                values[normal_mask],
                height=bar_height,
                color=COLORS[label],
                edgecolor="none",
                zorder=3
            )

        # Global 柱
        if np.any(global_mask):
            ax.barh(
                y_pos[global_mask] + offsets[label],
                values[global_mask],
                height=bar_height * 1.12,
                color=COLORS[label],
                edgecolor="black",
                linewidth=0.55,
                hatch="///",
                zorder=4
            )

    ax.axvline(
        0,
        color="black",
        linestyle="-",
        linewidth=0.50,
        alpha=0.75,
        zorder=2
    )

    ax.set_yticks(y_pos)
    tick_labels = ax.set_yticklabels(countries_plot_display)

    # Global 标签加粗
    for tick, is_g in zip(tick_labels, is_global_row):
        if is_g:
            tick.set_fontweight("bold")

    ax.invert_yaxis()

    x_values = plot_wide[DESIRED_COLS].values.astype(float)
    set_symmetric_xlim(ax, x_values, nbins=4)

    if show_ylabel:
        ax.set_ylabel("Global / Country")
    else:
        ax.set_ylabel("")

    style_axis(ax)

    ax.text(
        -0.10,
        1.03,
        panel_label,
        transform=ax.transAxes,
        fontsize=9,
        fontweight="bold",
        va="bottom",
        ha="left"
    )


# ============================================================
# 8. 主程序：同时计算 DFAI = 1 和 DFAI = -1
# ============================================================
country_regions = get_country_regions()

global_results = {}
stage_results = {}
country_results = {}
global_trend_results = {}
country_plot_results = {}

for dfai in DFAI_LIST:
    global_results[dfai] = {}
    stage_results[dfai] = {}
    country_results[dfai] = {}
    global_trend_results[dfai] = {}
    country_plot_results[dfai] = {}

    for crop in CROPS:
        global_df, stage_df, all_country_df, global_trend_df, plot_wide = process_crop(
            crop,
            dfai,
            country_regions
        )

        global_results[dfai][crop] = global_df
        stage_results[dfai][crop] = stage_df
        country_results[dfai][crop] = all_country_df
        global_trend_results[dfai][crop] = global_trend_df
        country_plot_results[dfai][crop] = plot_wide

# ============================================================
# 9. 合并保存所有 DFAI、所有作物的统计结果
# ============================================================
all_global_df = pd.concat(
    [
        global_results[dfai][crop]
        for dfai in DFAI_LIST
        for crop in CROPS
    ],
    ignore_index=True
)

valid_stage_dfs = [
    stage_results[dfai][crop]
    for dfai in DFAI_LIST
    for crop in CROPS
    if not stage_results[dfai][crop].empty
]

if len(valid_stage_dfs) > 0:
    all_stage_df = pd.concat(valid_stage_dfs, ignore_index=True)
else:
    all_stage_df = pd.DataFrame()

valid_country_dfs = [
    country_results[dfai][crop]
    for dfai in DFAI_LIST
    for crop in CROPS
    if not country_results[dfai][crop].empty
]

if len(valid_country_dfs) > 0:
    all_country_df = pd.concat(valid_country_dfs, ignore_index=True)

    # 保存总表前同样只保留有效栅格点数 > 10 的国家尺度记录。
    if "n_valid_grid_cells" in all_country_df.columns:
        all_country_df = all_country_df[
            all_country_df["n_valid_grid_cells"] > MIN_VALID_GRID_CELLS
        ].copy()
else:
    all_country_df = pd.DataFrame()

valid_global_trend_dfs = [
    global_trend_results[dfai][crop]
    for dfai in DFAI_LIST
    for crop in CROPS
    if not global_trend_results[dfai][crop].empty
]

if len(valid_global_trend_dfs) > 0:
    all_global_trend_df = pd.concat(valid_global_trend_dfs, ignore_index=True)
else:
    all_global_trend_df = pd.DataFrame()

all_global_csv = fr"{OUT_DIR}\all_crops_global_mean_residual_{START_YEAR}_{END_YEAR}_DFAI_1_minus1_{Data_from}.csv"
all_stage_csv = fr"{OUT_DIR}\all_crops_global_stage_trend_{START_YEAR}_{END_YEAR}_DFAI_1_minus1_{Data_from}.csv"
all_country_csv = fr"{OUT_DIR}\all_crops_country_trend_{START_YEAR}_{END_YEAR}_DFAI_1_minus1_{Data_from}.csv"
all_global_trend_csv = fr"{OUT_DIR}\all_crops_global_overall_trend_{START_YEAR}_{END_YEAR}_DFAI_1_minus1_{Data_from}.csv"

all_global_df.to_csv(all_global_csv, index=False, encoding="utf-8-sig")
all_stage_df.to_csv(all_stage_csv, index=False, encoding="utf-8-sig")
all_country_df.to_csv(all_country_csv, index=False, encoding="utf-8-sig")
all_global_trend_df.to_csv(all_global_trend_csv, index=False, encoding="utf-8-sig")

print("所有 DFAI、四作物全球均值总表已保存：", all_global_csv)
print("所有 DFAI、四作物全球阶段趋势总表已保存：", all_stage_csv)
print("所有 DFAI、四作物国家趋势总表已保存：", all_country_csv)
print("所有 DFAI、四作物全球整体趋势总表已保存：", all_global_trend_csv)

# ============================================================
# 9b. 将统计表导出为 Nature-like 三线表 Word
#     按作物 × DTF/FTD 分别生成 8 张表：
#     4 个作物 × 2 个 DFAI 类型 = 8 张表。
# ============================================================
all_tables_docx = fr"{OUT_DIR}\crop_DTF_FTD_separate_trend_tables_Nature_three_line_{START_YEAR}_{END_YEAR}_{Data_from}.docx"

eight_crop_dfai_tables = build_eight_crop_dfai_tables(
    all_global_trend_df=all_global_trend_df,
    all_country_df=all_country_df
)

save_tables_to_nature_word(
    tables=eight_crop_dfai_tables,
    out_docx=all_tables_docx,
    title=(
        f"Supplementary trend tables by crop and DTF/FTD "
        f"({START_YEAR}–{END_YEAR})"
    )
)

print("按 4 个作物 × DTF/FTD 分组的 8 张三线表已保存：", all_tables_docx)

# ============================================================
# 10. 生成 4 行 × 4 列组合图
#     第 1-2 行：DFAI = 1
#     第 3-4 行：DFAI = -1
# ============================================================
fig_width = 230 / 25.4
fig_height = 220 / 25.4

fig = plt.figure(figsize=(fig_width, fig_height))

gs = fig.add_gridspec(
    nrows=4,
    ncols=4,
    height_ratios=[1.0, 1.05, 1.0, 1.05],
    hspace=0.30,
    wspace=0.28
)

axes = np.empty((4, 4), dtype=object)

for r in range(4):
    for c in range(4):
        axes[r, c] = fig.add_subplot(gs[r, c])

panel_labels = list("abcdefghijklmnop")
panel_id = 0

for block_id, dfai in enumerate(DFAI_LIST):
    row_global = block_id * 2
    row_country = block_id * 2 + 1

    for i, crop in enumerate(CROPS):
        ax_g = axes[row_global, i]
        ax_c = axes[row_country, i]

        plot_global_panel(
            ax_g,
            crop,
            global_results[dfai][crop],
            panel_labels[panel_id],
            show_ylabel=(i == 0)
        )
        panel_id += 1

        plot_country_panel(
            ax_c,
            crop,
            country_plot_results[dfai][crop],
            panel_labels[panel_id],
            show_ylabel=(i == 0)
        )
        panel_id += 1

        # 只在最上面一行放作物标题，避免重复
        if row_global == 0:
            ax_g.set_title(
                CROP_TITLES[crop],
                fontsize=9,
                fontweight="bold",
                pad=4
            )

# ============================================================
# 11. 统一图例与说明
# ============================================================
legend_handles = [
    Line2D(
        [0],
        [0],
        color=COLORS[var],
        lw=1.2,
        label=LABELS[var]
    )
    for var in VARIABLES
]

fig.legend(
    handles=legend_handles,
    frameon=False,
    loc="upper center",
    bbox_to_anchor=(0.5, 0.992),
    ncol=3,
    handlelength=1.8,
    handletextpad=0.4,
    columnspacing=1.4
)

fig.text(
    0.5,
    0.968,
    "Solid lines: annual means; dashed lines: stage trends; shading: 95% CI",
    ha="center",
    va="center",
    fontsize=7
)

# 左侧标注 DFAI 分组
fig.text(
    0.012,
    0.735,
    "DTF",
    ha="center",
    va="center",
    rotation=90,
    fontsize=9,
    fontweight="bold"
)

fig.text(
    0.012,
    0.315,
    "FTD",
    ha="center",
    va="center",
    rotation=90,
    fontsize=9,
    fontweight="bold"
)

# Global + 国家趋势图统一 x 轴说明
fig.text(
    0.5,
    0.035,
    "Global and country-level trend in mean yield residual per decade, 1982–2016",
    ha="center",
    va="center",
    fontsize=8
)

fig.subplots_adjust(
    left=0.065,
    right=0.995,
    top=0.94,
    bottom=0.075
)

# ============================================================
# 12. 保存图片
# ============================================================
out_png = fr"{OUT_DIR}\all_crops_global_and_country_residual_4row4col_DFAI_1_minus1_{Data_from}.png"
out_pdf = fr"{OUT_DIR}\all_crops_global_and_country_residual_4row4col_DFAI_1_minus1_{Data_from}.pdf"

plt.savefig(out_png, dpi=600, bbox_inches="tight")
plt.savefig(out_pdf, bbox_inches="tight")
plt.close()

print("DFAI = 1 / -1 组合图 PNG 已保存：", out_png)
print("DFAI = 1 / -1 组合图 PDF 已保存：", out_pdf)
