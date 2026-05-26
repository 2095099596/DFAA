from pathlib import Path
import re

import numpy as np
import pandas as pd
import xarray as xr
import regionmask
import matplotlib.pyplot as plt

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    Document = None


# =========================
# 批量路径与参数设置
# =========================

# 数据源：会自动分别处理 ERA5、GPCC、CSIC
DATA_SOURCES = ["ERA5", "GPCC", "CSIC"]

# Num = 1 对应 DTF；Num = -1 对应 FTD
NUM_CONFIGS = {
    1: "DTF",
    -1: "FTD",
}

# 4 种农作物
CROPS = ["maize", "rice", "soybeans", "wheat"]

# 数据根目录模板
# 会自动替换为：
# C:\Users\lxh\Desktop\NongZuoWu_ChanLiang\New_result\Data\LDFAI_SPEI_ERA5_3_2
# C:\Users\lxh\Desktop\NongZuoWu_ChanLiang\New_result\Data\LDFAI_SPEI_GPCC_3_2
# C:\Users\lxh\Desktop\NongZuoWu_ChanLiang\New_result\Data\LDFAI_SPEI_CSIC_3_2
DATA_ROOT_TEMPLATE = (
    r"C:\Users\lxh\Desktop\NongZuoWu_ChanLiang\New_result\Data\LDFAI_SPEI_{data_from}_3_2"
)

# NC 文件模板
# Num = 1  -> ls_1.0_residual_mean_growing_...
# Num = -1 -> ls_-1.0_residual_mean_growing_...
NC_FILE_TEMPLATE = (
    r"multimodel_aic\ls_{num_token}_residual_mean_growing_-60_60_0_360.0s_multimodel_aic.nc"
)

output_dir = Path(
    r"D:\Python\DssAat\NO_2\Figure2a\Figure1c\Country"
)
output_dir.mkdir(parents=True, exist_ok=True)

# 只保留有效栅格数量大于该阈值的国家
MIN_GRID_COUNT = 1

# 如果国家太多导致图太密，可以改成 30、50 等
# None 表示绘制所有满足 MIN_GRID_COUNT 的国家
TOP_N = None

# 国家数量过多时，自动把横向柱状图拆成 2 或 3 列，避免单张图过长。
# 可选值："auto"、1、2、3。
# "auto" 会根据 MAX_COUNTRIES_PER_COLUMN 自动选择 1/2/3 列。
COUNTRY_PLOT_NCOLS = "auto"

# 每列最多显示多少个国家；超过后自动增加列数，最多 3 列。
# 如果仍然太高，可把 38 改小，例如 30；如果标签太密，可改大并配合 TOP_N。
MAX_COUNTRIES_PER_COLUMN = 38

# Word 三线表设置：None 表示全部写入；如果 Word 太大，可改为 100、200 等
WORD_TABLE_MAX_ROWS = None
WORD_FLOAT_DIGITS = 4

# Nature 风格配色：色盲友好，DTF/FTD 两类使用不同颜色
DIRECTION_COLORS = {
    "DTF": "#0072B2",  # blue
    "FTD": "#D55E00",  # vermillion
}

# 95% 误差线设置：使用 country_mean ± 1.96 × SEM。
# SEM 基于国家内有效栅格的面积加权标准差和 Kish effective sample size。
CI_LEVEL = 0.95
CI_Z_VALUE = 1.96
ERRORBAR_COLOR = "0.15"
ERRORBAR_LINEWIDTH = 0.45
ERRORBAR_CAPSIZE = 1.6
ERRORBAR_CAPTHICK = 0.45

# Word 表头浅色底纹，便于区分 DTF 与 FTD 列
DIRECTION_HEADER_FILL = {
    "DTF": "D9EAF7",
    "FTD": "F7E1D3",
}


# =========================
# Nature 风格绘图设置
# =========================

def apply_nature_style():
    """
    设置接近 Nature 期刊风格的 Matplotlib 参数：
    白底、无网格、细坐标轴、小字号、高分辨率、PDF/PS 保留可编辑文字。
    """
    plt.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 7,
        "axes.labelsize": 7,
        "axes.titlesize": 8,
        "xtick.labelsize": 6.5,
        "ytick.labelsize": 6.5,
        "legend.fontsize": 7,
        "axes.linewidth": 0.5,
        "axes.grid": False,
        "xtick.major.width": 0.5,
        "ytick.major.width": 0.5,
        "xtick.major.size": 2.5,
        "ytick.major.size": 2.5,
        "xtick.direction": "out",
        "ytick.direction": "out",
        "figure.dpi": 150,
        "savefig.dpi": 600,
        "savefig.transparent": False,
        "savefig.facecolor": "white",
        "pdf.fonttype": 42,
        "ps.fonttype": 42,
        "mathtext.default": "regular",
        "axes.unicode_minus": False,
    })

def mm_to_inch(mm):
    return mm / 25.4


# =========================
# 基础函数
# =========================

def get_data_root(data_from):
    return Path(DATA_ROOT_TEMPLATE.format(data_from=data_from))


def get_nc_file(num):
    num_token = f"{float(num):.1f}"
    return NC_FILE_TEMPLATE.format(num_token=num_token)


def find_nc_path(data_from, crop, num):
    """
    优先寻找 DATA_ROOT/crop/NC_FILE；
    如果没有，再寻找 DATA_ROOT/NC_FILE。
    """
    data_root = get_data_root(data_from)
    nc_file = get_nc_file(num)

    path1 = data_root / crop / nc_file
    path2 = data_root / nc_file

    if path1.exists():
        return path1
    if path2.exists():
        return path2

    return None


def find_coord(ds, names):
    for name in names:
        if name in ds.coords or name in ds.dims:
            return name
    raise ValueError(f"找不到坐标变量：{names}")


def get_country_regions():
    """
    获取 Natural Earth 国家边界，兼容不同版本 regionmask。
    """
    try:
        return regionmask.defined_regions.natural_earth_v5_0_0.countries_110
    except AttributeError:
        return regionmask.defined_regions.natural_earth.countries_110


def standardize_lon(ds, lon_name):
    """
    将 0–360 经度转换为 -180–180，便于和国家边界匹配。
    """
    lon = ds[lon_name]

    if float(lon.max()) > 180:
        new_lon = ((lon + 180) % 360) - 180
        ds = ds.assign_coords({lon_name: new_lon})
        ds = ds.sortby(lon_name)

    return ds


def safe_filename(text):
    """
    防止变量名中出现 Windows 文件名非法字符。
    """
    text = str(text)
    return re.sub(r'[\\/:*?"<>|]+', "_", text)


# =========================
# 计算每个国家平均值
# =========================

def calc_country_mean_for_ds(ds, crop_name, data_from, num):
    """
    对单个 NC 文件计算各国家面积加权平均值，并统计有效栅格数量。

    只保留 grid_count > MIN_GRID_COUNT 的国家。
    """
    lat_name = find_coord(ds, ["lat", "latitude", "y"])
    lon_name = find_coord(ds, ["lon", "longitude", "x"])

    ds = standardize_lon(ds, lon_name)

    countries = get_country_regions()

    # 生成国家掩膜：region × lat × lon
    mask_3d = countries.mask_3D(ds[lon_name], ds[lat_name])

    # 纬度面积权重
    weights = np.cos(np.deg2rad(ds[lat_name]))

    # 国家编号、名称、缩写映射
    num_to_name = dict(zip(countries.numbers, countries.names))
    num_to_abbrev = dict(zip(countries.numbers, countries.abbrevs))

    # 找出包含 lat/lon 的数值变量
    spatial_vars = []
    for var in ds.data_vars:
        da = ds[var]

        if (
            lat_name in da.dims
            and lon_name in da.dims
            and np.issubdtype(da.dtype, np.number)
        ):
            spatial_vars.append(var)

    if not spatial_vars:
        raise ValueError("NC 文件中没有找到包含 lat/lon 维度的数值变量。")

    results = []

    for var in spatial_vars:
        print(f"  正在计算变量：{var}")

        da = ds[var]

        # 有效栅格：属于该国家，且变量不是 NaN
        valid_mask = mask_3d & da.notnull()

        # 统计每个国家参与计算的有效栅格数量
        grid_count = valid_mask.sum(dim=[lat_name, lon_name], skipna=True)

        # 面积加权国家平均：
        # mean = sum(value * cos(lat)) / sum(cos(lat))
        numerator = (
            (da * weights)
            .where(mask_3d)
            .sum(dim=[lat_name, lon_name], skipna=True)
        )

        denominator = (
            weights
            .where(valid_mask)
            .sum(dim=[lat_name, lon_name], skipna=True)
        )

        country_mean = numerator / denominator

        # 95% 误差线：国家内有效栅格的面积加权均值 ± 1.96 × SEM。
        # 由于使用纬度面积权重，SEM 中的样本量采用 Kish effective sample size：
        # n_eff = (sum(w))^2 / sum(w^2)。
        # 当国家内有效格点过少时，误差线设为 NaN，避免误导。
        weight_sum = denominator
        weight_sq_sum = (
            (weights ** 2)
            .where(valid_mask)
            .sum(dim=[lat_name, lon_name], skipna=True)
        )
        country_neff = (weight_sum ** 2) / weight_sq_sum

        weighted_variance = (
            (((da - country_mean) ** 2) * weights)
            .where(valid_mask)
            .sum(dim=[lat_name, lon_name], skipna=True)
            / weight_sum
        )
        country_std = np.sqrt(weighted_variance)
        country_sem = country_std / np.sqrt(country_neff)
        country_ci95 = CI_Z_VALUE * country_sem
        country_ci95 = country_ci95.where(country_neff > 1)

        stats_ds = xr.Dataset(
            {
                "country_mean": country_mean,
                "grid_count": grid_count,
                "country_neff": country_neff,
                "country_sem": country_sem,
                "country_ci95": country_ci95,
            }
        )
        df = stats_ds.to_dataframe().reset_index()

        df["country"] = df["region"].apply(
            lambda x: num_to_name.get(int(x), None)
        )
        df["iso3"] = df["region"].apply(
            lambda x: num_to_abbrev.get(int(x), None)
        )

        df["crop"] = crop_name
        df["data_from"] = data_from
        df["num"] = num
        df["direction"] = NUM_CONFIGS[num]
        df["variable"] = var

        df = df.dropna(subset=["country_mean"])
        df["grid_count"] = df["grid_count"].fillna(0).astype(int)

        before_filter = len(df)
        df = df[df["grid_count"] > MIN_GRID_COUNT].copy()
        after_filter = len(df)

        print(
            f"  {data_from} | {crop_name} | {NUM_CONFIGS[num]} | {var}: "
            f"保留 grid_count > {MIN_GRID_COUNT} 的记录 {after_filter}/{before_filter}"
        )

        # 调整列顺序
        first_cols = [
            "crop",
            "data_from",
            "num",
            "direction",
            "variable",
            "country",
            "iso3",
            "region",
            "grid_count",
            "country_mean",
            "country_sem",
            "country_ci95",
            "country_neff",
        ]
        other_cols = [c for c in df.columns if c not in first_cols]
        df = df[first_cols + other_cols]

        results.append(df)

    if not results:
        return pd.DataFrame()

    return pd.concat(results, ignore_index=True)


# =========================
# 合并 DTF 与 FTD 表格
# =========================

def make_combined_dtf_ftd_table(df_subset):
    """
    将同一个 crop × data_from 的 DTF 和 FTD 合并到同一张表。

    输出宽表列：
    crop, data_from, variable, country, iso3, region,
    DTF_mean, FTD_mean, DTF_ci95, FTD_ci95,
    DTF_sem, FTD_sem, DTF_neff, FTD_neff,
    DTF_grid_count, FTD_grid_count
    """
    if df_subset.empty:
        return pd.DataFrame()

    id_cols = ["crop", "data_from", "variable", "country", "iso3", "region"]

    # 如果原始数据还有额外维度，此处对同一国家、同一方向先取平均，
    # 保证 DTF 和 FTD 在同一行便于出图和制表。
    grouped = (
        df_subset
        .groupby(id_cols + ["direction"], as_index=False)
        .agg(
            country_mean=("country_mean", "mean"),
            country_ci95=("country_ci95", "mean"),
            country_sem=("country_sem", "mean"),
            country_neff=("country_neff", "mean"),
            grid_count=("grid_count", "max"),
        )
    )

    mean_wide = (
        grouped
        .pivot_table(
            index=id_cols,
            columns="direction",
            values="country_mean",
            aggfunc="mean",
        )
        .reset_index()
    )

    count_wide = (
        grouped
        .pivot_table(
            index=id_cols,
            columns="direction",
            values="grid_count",
            aggfunc="max",
        )
        .reset_index()
    )

    ci95_wide = (
        grouped
        .pivot_table(
            index=id_cols,
            columns="direction",
            values="country_ci95",
            aggfunc="mean",
        )
        .reset_index()
    )

    sem_wide = (
        grouped
        .pivot_table(
            index=id_cols,
            columns="direction",
            values="country_sem",
            aggfunc="mean",
        )
        .reset_index()
    )

    neff_wide = (
        grouped
        .pivot_table(
            index=id_cols,
            columns="direction",
            values="country_neff",
            aggfunc="mean",
        )
        .reset_index()
    )

    mean_wide = mean_wide.rename(
        columns={
            "DTF": "DTF_mean",
            "FTD": "FTD_mean",
        }
    )
    count_wide = count_wide.rename(
        columns={
            "DTF": "DTF_grid_count",
            "FTD": "FTD_grid_count",
        }
    )
    ci95_wide = ci95_wide.rename(
        columns={
            "DTF": "DTF_ci95",
            "FTD": "FTD_ci95",
        }
    )
    sem_wide = sem_wide.rename(
        columns={
            "DTF": "DTF_sem",
            "FTD": "FTD_sem",
        }
    )
    neff_wide = neff_wide.rename(
        columns={
            "DTF": "DTF_neff",
            "FTD": "FTD_neff",
        }
    )

    table_df = mean_wide.merge(count_wide, on=id_cols, how="left")
    table_df = table_df.merge(ci95_wide, on=id_cols, how="left")
    table_df = table_df.merge(sem_wide, on=id_cols, how="left")
    table_df = table_df.merge(neff_wide, on=id_cols, how="left")

    expected_cols = [
        "crop",
        "data_from",
        "variable",
        "country",
        "iso3",
        "region",
        "DTF_mean",
        "FTD_mean",
        "DTF_ci95",
        "FTD_ci95",
        "DTF_sem",
        "FTD_sem",
        "DTF_neff",
        "FTD_neff",
        "DTF_grid_count",
        "FTD_grid_count",
    ]
    for col in expected_cols:
        if col not in table_df.columns:
            table_df[col] = np.nan

    table_df = table_df[expected_cols]

    # 按 DTF 和 FTD 平均绝对值排序，便于查看重要国家
    sort_value = table_df[["DTF_mean", "FTD_mean"]].abs().max(axis=1)
    table_df = (
        table_df
        .assign(_sort_value=sort_value)
        .sort_values("_sort_value", ascending=False)
        .drop(columns="_sort_value")
        .reset_index(drop=True)
    )

    # grid_count 列转为整数型的可空类型
    for col in ["DTF_grid_count", "FTD_grid_count"]:
        table_df[col] = table_df[col].round().astype("Int64")

    return table_df


# =========================
# Word 三线表输出
# =========================

def _ensure_python_docx():
    if Document is None:
        raise ImportError(
            "缺少 python-docx。请先在当前 Python 环境中安装：pip install python-docx"
        )


def _set_cell_border(cell, **kwargs):
    """
    设置单元格边框。
    示例：_set_cell_border(cell, top={"val": "single", "sz": "8", "color": "000000"})
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue

        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)

        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def _clear_cell_borders(cell):
    nil_border = {"val": "nil"}
    _set_cell_border(
        cell,
        top=nil_border,
        bottom=nil_border,
        left=nil_border,
        right=nil_border,
        insideH=nil_border,
        insideV=nil_border,
    )


def _set_cell_shading(cell, fill):
    """
    设置单元格底纹颜色。fill 为 6 位十六进制颜色，不带 #。
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text, bold=False, font_size=8):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(font_size)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _format_table_value(value, digits=4):
    if pd.isna(value):
        return ""
    if isinstance(value, (np.integer, int)):
        return str(int(value))
    if isinstance(value, (np.floating, float)):
        return f"{float(value):.{digits}g}"
    return str(value)


def save_dataframe_to_word_three_line_table(
    df,
    out_docx,
    title="Table 1. Country mean by crop",
    max_rows=None,
    float_digits=4,
):
    """
    将 DataFrame 保存为 Word 三线表：
    顶线、表头下线、底线；无竖线和内部横线。

    DTF 与 FTD 相关列的表头会使用浅色底纹，以对应图中的两种颜色。
    """
    _ensure_python_docx()

    out_docx = Path(out_docx)
    out_docx.parent.mkdir(parents=True, exist_ok=True)

    table_df = df.copy()
    if max_rows is not None:
        table_df = table_df.head(max_rows)

    doc = Document()

    # 横向页面，适合列数较多的表格
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)

    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 写表头
    header_cells = table.rows[0].cells
    for j, col in enumerate(table_df.columns):
        _set_cell_text(header_cells[j], col, bold=True, font_size=8)

        if str(col).startswith("DTF"):
            _set_cell_shading(header_cells[j], DIRECTION_HEADER_FILL["DTF"])
        elif str(col).startswith("FTD"):
            _set_cell_shading(header_cells[j], DIRECTION_HEADER_FILL["FTD"])

    # 写数据行
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(table_df.columns):
            text = _format_table_value(row[col], digits=float_digits)
            _set_cell_text(cells[j], text, bold=False, font_size=7)

    # 清除所有边框
    for row in table.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)

    # 三线表：顶线、表头下线、底线
    top_line = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    mid_line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    bottom_line = {"val": "single", "sz": "12", "space": "0", "color": "000000"}

    for cell in table.rows[0].cells:
        _set_cell_border(cell, top=top_line, bottom=mid_line)

    if len(table.rows) > 0:
        for cell in table.rows[-1].cells:
            _set_cell_border(cell, bottom=bottom_line)

    if max_rows is not None and len(df) > max_rows:
        note = doc.add_paragraph(
            f"Note: The table shows the first {max_rows} rows out of {len(df)} rows. "
            f"Only countries with grid_count > {MIN_GRID_COUNT} are included."
        )
        note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in note.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(8)

    doc.save(out_docx)
    print(f"Word 三线表已保存：{out_docx}")


# =========================
# 绘制 DTF 与 FTD 合并 Nature 风格柱状图
# =========================

def _choose_country_plot_ncols(n_items):
    """
    根据国家数量选择绘图列数。
    - COUNTRY_PLOT_NCOLS = 1/2/3 时固定列数；
    - COUNTRY_PLOT_NCOLS = "auto" 时，根据 MAX_COUNTRIES_PER_COLUMN 自动选择。
    """
    if n_items <= 0:
        return 1

    if COUNTRY_PLOT_NCOLS in [1, 2, 3]:
        return int(COUNTRY_PLOT_NCOLS)

    if n_items <= MAX_COUNTRIES_PER_COLUMN:
        return 1
    if n_items <= MAX_COUNTRIES_PER_COLUMN * 2:
        return 2
    return 3


def _get_symmetric_xlim(*arrays, pad=0.08):
    """
    获取以 0 为中心的对称 x 轴范围，便于比较 DTF 与 FTD。
    """
    values = []
    for arr in arrays:
        arr = np.asarray(arr, dtype=float)
        arr = arr[~np.isnan(arr)]
        if arr.size > 0:
            values.append(arr)

    if not values:
        return -1.0, 1.0

    values = np.concatenate(values)
    max_abs = np.nanmax(np.abs(values))
    if not np.isfinite(max_abs) or max_abs == 0:
        max_abs = 1.0
    max_abs = max_abs * (1 + pad)
    return -max_abs, max_abs


def _plot_one_country_column(ax, plot_df, xlim, show_xlabel=False):
    """
    在一个子图列中绘制 DTF/FTD 横向柱状图。
    """
    y = np.arange(len(plot_df))
    bar_height = 0.34

    dtf_values = plot_df["DTF_mean"].to_numpy(dtype=float)
    ftd_values = plot_df["FTD_mean"].to_numpy(dtype=float)

    dtf_ci95 = (
        plot_df["DTF_ci95"].to_numpy(dtype=float)
        if "DTF_ci95" in plot_df.columns
        else np.full(len(plot_df), np.nan)
    )
    ftd_ci95 = (
        plot_df["FTD_ci95"].to_numpy(dtype=float)
        if "FTD_ci95" in plot_df.columns
        else np.full(len(plot_df), np.nan)
    )

    dtf_mask = ~np.isnan(dtf_values)
    ftd_mask = ~np.isnan(ftd_values)
    dtf_err_mask = dtf_mask & np.isfinite(dtf_ci95)
    ftd_err_mask = ftd_mask & np.isfinite(ftd_ci95)

    ax.barh(
        y[dtf_mask] - bar_height / 2,
        dtf_values[dtf_mask],
        height=bar_height,
        color=DIRECTION_COLORS["DTF"],
        edgecolor="none",
        label="DTF (Num = 1)",
        rasterized=False,
    )

    ax.barh(
        y[ftd_mask] + bar_height / 2,
        ftd_values[ftd_mask],
        height=bar_height,
        color=DIRECTION_COLORS["FTD"],
        edgecolor="none",
        label="FTD (Num = -1)",
        rasterized=False,
    )

    # 95% 误差线：横向柱状图使用 xerr。
    # fmt="none" 表示只画误差线，不额外叠加点。
    if np.any(dtf_err_mask):
        ax.errorbar(
            dtf_values[dtf_err_mask],
            y[dtf_err_mask] - bar_height / 2,
            xerr=dtf_ci95[dtf_err_mask],
            fmt="none",
            ecolor=ERRORBAR_COLOR,
            elinewidth=ERRORBAR_LINEWIDTH,
            capsize=ERRORBAR_CAPSIZE,
            capthick=ERRORBAR_CAPTHICK,
            zorder=3,
        )

    if np.any(ftd_err_mask):
        ax.errorbar(
            ftd_values[ftd_err_mask],
            y[ftd_err_mask] + bar_height / 2,
            xerr=ftd_ci95[ftd_err_mask],
            fmt="none",
            ecolor=ERRORBAR_COLOR,
            elinewidth=ERRORBAR_LINEWIDTH,
            capsize=ERRORBAR_CAPSIZE,
            capthick=ERRORBAR_CAPTHICK,
            zorder=3,
        )

    labels = plot_df["country"].astype(str).tolist()
    ax.set_yticks(y)
    ax.set_yticklabels(labels)
    ax.invert_yaxis()

    ax.set_xlim(*xlim)
    ax.axvline(0, linewidth=0.5, color="0.20", zorder=0)

    if show_xlabel:
        ax.set_xlabel("Country mean ± 95% CI")
    else:
        ax.set_xlabel("")
    ax.set_ylabel("")

    # Nature 风格：去掉上、右边框，保留细坐标轴。
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_linewidth(0.5)
    ax.spines["bottom"].set_linewidth(0.5)
    ax.tick_params(axis="both", width=0.5, length=2.5, pad=1.5)
    ax.grid(False)


def plot_combined_dtf_ftd_bar_chart(table_df, crop, data_from, variable, output_dir, top_n=None):
    """
    每个 crop × data_from × variable 生成一张 Nature 风格图：
    同一个国家显示 DTF 与 FTD 两个横向柱。

    改进点：
    1. 图宽采用 Nature 常用版面宽度：单列约 89 mm，双栏约 183 mm；
    2. 国家数量过多时自动拆分为 2 或 3 个并排子图列，避免单张图片过长；
    3. 所有子图共享同一个以 0 为中心的 x 轴范围，便于跨列比较；
    4. 柱体上叠加 95% 误差线（country_mean ± 1.96 × SEM）。
    """
    if table_df.empty:
        print(f"{crop} | {data_from} | {variable} 没有可绘图数据")
        return

    apply_nature_style()

    plot_dir = output_dir / "combined_DTF_FTD_country_bar_charts_nature"
    plot_dir.mkdir(parents=True, exist_ok=True)

    plot_df = table_df.copy()

    # 如果只想画前 N 个国家，按 DTF/FTD 最大绝对值筛选。
    if top_n is not None:
        plot_df = (
            plot_df
            .assign(abs_value=plot_df[["DTF_mean", "FTD_mean"]].abs().max(axis=1))
            .sort_values("abs_value", ascending=False)
            .head(top_n)
            .drop(columns="abs_value")
        )

    # 按 DTF/FTD 平均值排序，方便观察。
    plot_df = (
        plot_df
        .assign(sort_value=plot_df[["DTF_mean", "FTD_mean"]].mean(axis=1, skipna=True))
        .sort_values("sort_value", ascending=True)
        .drop(columns="sort_value")
        .reset_index(drop=True)
    )

    n_items = len(plot_df)
    ncols = _choose_country_plot_ncols(n_items)
    ncols = min(ncols, n_items) if n_items > 0 else 1

    # 按顺序切成多列；第一列、第二列、第三列从左到右依次接续。
    index_chunks = np.array_split(np.arange(n_items), ncols)
    df_chunks = [plot_df.iloc[idx].reset_index(drop=True) for idx in index_chunks if len(idx) > 0]
    ncols = len(df_chunks)
    max_rows_per_col = max(len(df_col) for df_col in df_chunks)

    # 单列用 Nature 单栏宽度；多列用双栏宽度。
    fig_width = mm_to_inch(89 if ncols == 1 else 183)
    # 多列后高度按每列最大行数计算，避免原来 len(plot_df) * 0.24 造成图片过长。
    fig_height = max(mm_to_inch(70), max_rows_per_col * 0.20 + 0.75)

    fig, axes = plt.subplots(
        1,
        ncols,
        figsize=(fig_width, fig_height),
        sharex=True,
        squeeze=False,
        constrained_layout=False,
    )
    axes = axes.ravel()

    # x 轴范围同时考虑误差线，防止 95% CI 被裁切。
    dtf_xmax = plot_df["DTF_mean"] + plot_df.get("DTF_ci95", 0)
    dtf_xmin = plot_df["DTF_mean"] - plot_df.get("DTF_ci95", 0)
    ftd_xmax = plot_df["FTD_mean"] + plot_df.get("FTD_ci95", 0)
    ftd_xmin = plot_df["FTD_mean"] - plot_df.get("FTD_ci95", 0)
    xlim = _get_symmetric_xlim(dtf_xmax, dtf_xmin, ftd_xmax, ftd_xmin)

    legend_handles = None
    legend_labels = None

    for i, (ax, df_col) in enumerate(zip(axes, df_chunks)):
        _plot_one_country_column(
            ax=ax,
            plot_df=df_col,
            xlim=xlim,
            show_xlabel=True,
        )
        if legend_handles is None:
            legend_handles, legend_labels = ax.get_legend_handles_labels()

        # 多列时缩短每个子图的 x 轴刻度数量，减少拥挤。
        ax.locator_params(axis="x", nbins=4 if ncols >= 3 else 5)

    # 统一标题和图例，避免每列重复。
    fig.suptitle(f"{crop} | {data_from} | {variable}", y=0.995, fontsize=8, fontweight="normal")
    if legend_handles:
        fig.legend(
            legend_handles,
            legend_labels,
            frameon=False,
            loc="upper center",
            bbox_to_anchor=(0.5, 0.965),
            ncol=2,
            handlelength=1.2,
            columnspacing=1.2,
            borderaxespad=0,
        )

    # 为标题和图例留出上边距；多列时适当增大列间距。
    fig.subplots_adjust(
        left=0.10 if ncols == 1 else 0.06,
        right=0.99,
        bottom=0.07,
        top=0.90,
        wspace=0.45 if ncols >= 3 else 0.35,
    )

    base_name = (
        f"{safe_filename(crop)}_{safe_filename(data_from)}_"
        f"{safe_filename(variable)}_DTF_FTD_country_mean_95CI_nature_{ncols}col"
    )
    out_png = plot_dir / f"{base_name}.png"
    out_pdf = plot_dir / f"{base_name}.pdf"

    fig.savefig(out_png, dpi=600, bbox_inches="tight")
    fig.savefig(out_pdf, bbox_inches="tight")
    plt.close(fig)

    print(f"合并 DTF/FTD Nature 风格 {ncols} 列图已保存：{out_png}")
    print(f"合并 DTF/FTD Nature 风格 {ncols} 列矢量图已保存：{out_pdf}")


# =========================
# 输出每个 crop × data_from 的图和表
# =========================

def save_combined_tables_and_plots(final_df):
    """
    输出：
    1. 每个 crop × data_from 一个 CSV 宽表；
    2. 每个 crop × data_from 一个 Word 三线表；
    3. 每个 crop × data_from × variable 一个合并 DTF/FTD 图。

    如果每个 NC 只有一个变量，则图像数量为：
    4 个作物 × 3 个数据源 = 12 张。
    表格数量也为：
    4 个作物 × 3 个数据源 = 12 个 CSV + 12 个 Word。
    """
    table_dir = output_dir / "combined_DTF_FTD_tables"
    table_dir.mkdir(parents=True, exist_ok=True)

    all_wide_tables = []

    for (crop, data_from), df_cd in final_df.groupby(["crop", "data_from"]):
        table_df = make_combined_dtf_ftd_table(df_cd)

        if table_df.empty:
            print(f"{crop} | {data_from} 没有可输出表格数据")
            continue

        all_wide_tables.append(table_df)

        out_csv = table_dir / f"country_mean_{safe_filename(crop)}_{safe_filename(data_from)}_DTF_FTD_gridgt{MIN_GRID_COUNT}.csv"
        table_df.to_csv(out_csv, index=False, encoding="utf-8-sig")
        print(f"合并 DTF/FTD CSV 表格已保存：{out_csv}")

        out_docx = table_dir / f"country_mean_{safe_filename(crop)}_{safe_filename(data_from)}_DTF_FTD_gridgt{MIN_GRID_COUNT}_three_line_table.docx"
        save_dataframe_to_word_three_line_table(
            df=table_df,
            out_docx=out_docx,
            title=f"Table. Country mean for {crop} using {data_from}: DTF vs FTD",
            max_rows=WORD_TABLE_MAX_ROWS,
            float_digits=WORD_FLOAT_DIGITS,
        )

        # 画图：一般每个 crop × data_from 只有一个 variable，因此生成 12 张图。
        # 如果存在多个 variable，则每个 variable 额外生成一张图。
        for variable, table_var in table_df.groupby("variable"):
            plot_combined_dtf_ftd_bar_chart(
                table_df=table_var,
                crop=crop,
                data_from=data_from,
                variable=variable,
                output_dir=output_dir,
                top_n=TOP_N,
            )

    if all_wide_tables:
        all_wide_df = pd.concat(all_wide_tables, ignore_index=True)
        out_all_wide = table_dir / f"country_mean_all_sources_DTF_FTD_gridgt{MIN_GRID_COUNT}_wide.csv"
        all_wide_df.to_csv(out_all_wide, index=False, encoding="utf-8-sig")
        print(f"总宽表已保存：{out_all_wide}")


# =========================
# 主程序
# =========================

def main():
    all_results = []

    for data_from in DATA_SOURCES:
        for crop in CROPS:
            for num, direction in NUM_CONFIGS.items():

                nc_path = find_nc_path(data_from=data_from, crop=crop, num=num)

                if nc_path is None:
                    print("=" * 80)
                    print(
                        f"找不到文件：data_from={data_from}, crop={crop}, "
                        f"Num={num} ({direction})"
                    )
                    print(f"已尝试：{get_data_root(data_from) / crop / get_nc_file(num)}")
                    print(f"已尝试：{get_data_root(data_from) / get_nc_file(num)}")
                    continue

                print("=" * 80)
                print(f"正在处理数据源：{data_from}")
                print(f"正在处理作物：{crop}")
                print(f"正在处理方向：{direction}，Num = {num}")
                print(f"NC 文件：{nc_path}")

                ds = xr.open_dataset(nc_path)

                try:
                    df_one = calc_country_mean_for_ds(
                        ds=ds,
                        crop_name=crop,
                        data_from=data_from,
                        num=num,
                    )
                    if not df_one.empty:
                        all_results.append(df_one)
                finally:
                    ds.close()

    if all_results:
        final_df = pd.concat(all_results, ignore_index=True)

        # 保存所有数据的长表，包含 data_from、crop、Num、direction、grid_count、country_mean、country_ci95
        out_all_long = output_dir / f"country_mean_all_sources_DTF_FTD_gridgt{MIN_GRID_COUNT}_long.csv"
        final_df.to_csv(out_all_long, index=False, encoding="utf-8-sig")

        print("=" * 80)
        print(f"所有数据源与方向的长表已保存：{out_all_long}")

        # 输出每个 crop × data_from 的合并表格和合并图像
        save_combined_tables_and_plots(final_df)

        print("=" * 80)
        print("全部完成。")

    else:
        print("没有成功处理任何文件。")


if __name__ == "__main__":
    main()
